#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch

# Criar planilha XLSX com múltiplas abas
print("Criando planilha.xlsx...")
wb = Workbook()

# Aba 1: Dados
ws1 = wb.active
ws1.title = "Dados"
ws1['A1'] = "Nome"
ws1['B1'] = "Idade"
ws1['C1'] = "Cidade"
ws1['A1'].font = Font(bold=True)
ws1['B1'].font = Font(bold=True)
ws1['C1'].font = Font(bold=True)

dados = [
    ["João Silva", 28, "São Paulo"],
    ["Maria Santos", 32, "Rio de Janeiro"],
    ["Pedro Oliveira", 25, "Belo Horizonte"],
    ["Ana Costa", 30, "Curitiba"],
]

for i, linha in enumerate(dados, start=2):
    ws1[f'A{i}'] = linha[0]
    ws1[f'B{i}'] = linha[1]
    ws1[f'C{i}'] = linha[2]

# Aba 2: Vendas
ws2 = wb.create_sheet("Vendas")
ws2['A1'] = "Produto"
ws2['B1'] = "Quantidade"
ws2['C1'] = "Valor"
ws2['A1'].font = Font(bold=True)
ws2['B1'].font = Font(bold=True)
ws2['C1'].font = Font(bold=True)

vendas = [
    ["Produto A", 100, 1500.00],
    ["Produto B", 50, 2500.00],
    ["Produto C", 75, 1800.00],
]

for i, linha in enumerate(vendas, start=2):
    ws2[f'A{i}'] = linha[0]
    ws2[f'B{i}'] = linha[1]
    ws2[f'C{i}'] = linha[2]

# Aba 3: Resumo
ws3 = wb.create_sheet("Resumo")
ws3['A1'] = "Resumo Geral"
ws3['A1'].font = Font(bold=True, size=14)
ws3['A3'] = "Total de pessoas:"
ws3['B3'] = len(dados)
ws3['A4'] = "Total de produtos:"
ws3['B4'] = len(vendas)

wb.save("planilha.xlsx")
print("✓ planilha.xlsx criada")

# Criar documento DOCX
print("Criando documento.docx...")
doc = Document()

# Título
doc.add_heading('Documento de Teste', 0)

# Parágrafo
doc.add_paragraph('Este é um documento Word de teste criado com python-docx.')

# Lista
doc.add_heading('Lista de Itens:', level=2)
doc.add_paragraph('Primeiro item', style='List Bullet')
doc.add_paragraph('Segundo item', style='List Bullet')
doc.add_paragraph('Terceiro item', style='List Bullet')

# Tabela
doc.add_heading('Tabela de Exemplo:', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'

# Cabeçalho
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Nome'
hdr_cells[1].text = 'Idade'
hdr_cells[2].text = 'Cidade'

# Dados
data = [
    ['João Silva', '28', 'São Paulo'],
    ['Maria Santos', '32', 'Rio de Janeiro'],
    ['Pedro Oliveira', '25', 'Belo Horizonte']
]

for i, row_data in enumerate(data, start=1):
    row_cells = table.rows[i].cells
    for j, value in enumerate(row_data):
        row_cells[j].text = value

doc.save('documento.docx')
print("✓ documento.docx criado")

# Criar documento DOC legado (será DOCX com nome .doc)
doc2 = Document()
doc2.add_heading('Documento Legado', 0)
doc2.add_paragraph('Este é um documento de formato legado (será salvo como .docx mas com extensão .doc)')
doc2.save('documento.doc')
print("✓ documento.doc criado")

# Criar PDF
print("Criando relatorio.pdf...")
pdf = SimpleDocTemplate("relatorio.pdf", pagesize=letter)
story = []
styles = getSampleStyleSheet()

# Título
title = Paragraph("Relatório de Teste", styles['Title'])
story.append(title)
story.append(Spacer(1, 12))

# Conteúdo
content = Paragraph(
    "Este é um relatório em PDF criado com ReportLab. " 
    "Ele contém texto formatado, tabelas e outros elementos.",
    styles['BodyText']
)
story.append(content)
story.append(Spacer(1, 12))

# Subtítulo
subtitle = Paragraph("Dados da Tabela", styles['Heading2'])
story.append(subtitle)
story.append(Spacer(1, 12))

# Tabela
data_table = [
    ['Nome', 'Idade', 'Cidade'],
    ['João Silva', '28', 'São Paulo'],
    ['Maria Santos', '32', 'Rio de Janeiro'],
    ['Pedro Oliveira', '25', 'Belo Horizonte'],
]

table = Table(data_table)
table.setStyle(TableStyle([
    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
    ('FONTSIZE', (0, 0), (-1, 0), 12),
    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
    ('GRID', (0, 0), (-1, -1), 1, colors.black)
]))

story.append(table)
pdf.build(story)
print("✓ relatorio.pdf criado")

print("\n✅ Todos os documentos foram criados com sucesso!")
